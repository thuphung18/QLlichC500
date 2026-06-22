import os
import sys
from fastapi.testclient import TestClient
import docx

# Add current directory to path
sys.path.append(os.path.abspath(os.path.dirname(__file__) + "/.."))

from main import app
from core.security import create_access_token

client = TestClient(app)

# Generate a valid access token for user "u001" (Admin)
token = create_access_token(data={"sub": "u001", "role": "admin"})

headers = {
    "Authorization": f"Bearer {token}"
}

# Create a real docx file using python-docx
doc = docx.Document()
doc.add_heading("LỊCH CÔNG TÁC TUẦN", level=1)
doc.add_paragraph("Tuần từ 22/06/2026 đến 28/06/2026")

# Add a table with schedule details
table = doc.add_table(rows=3, cols=5)
# Set header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Thứ/Ngày"
hdr_cells[1].text = "Thời gian"
hdr_cells[2].text = "Nội dung công việc"
hdr_cells[3].text = "Chủ trì"
hdr_cells[4].text = "Địa điểm"

# Row 1
row_cells = table.rows[1].cells
row_cells[0].text = "Thứ Hai\n22/06"
row_cells[1].text = "08:00 - 11:30"
row_cells[2].text = "Họp Giao ban đầu tuần phòng Quản lý đào tạo (QLĐT)"
row_cells[3].text = "Đ/c Vũ (PGĐ)"
row_cells[4].text = "Phòng hội trường lớn"

# Row 2
row_cells = table.rows[2].cells
row_cells[0].text = "Thứ Ba\n23/06"
row_cells[1].text = "14:00 - 17:00"
row_cells[2].text = "Tập huấn sử dụng phần mềm QL Lịch tuần khoa Hành chính (HC)"
row_cells[3].text = "Đ/c Nam (Trưởng phòng)"
row_cells[4].text = "Phòng họp tầng 3"

# Save the document
temp_doc_path = "temp_test_schedule.docx"
doc.save(temp_doc_path)

try:
    with open(temp_doc_path, "rb") as f:
        file_content = f.read()
        
    response = client.post(
        "/api/schedules/import",
        headers=headers,
        files={"file": ("test.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    print("Response status:", response.status_code)
    print("Response JSON:")
    import json
    print(json.dumps(response.json(), indent=2, ensure_ascii=False))
except Exception as e:
    print("Test failed with exception:", e)
finally:
    # Clean up temp file
    if os.path.exists(temp_doc_path):
        os.remove(temp_doc_path)
