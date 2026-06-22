import os
import sys
import json
import urllib.request
import ssl

# Create a test docx file
import docx
doc = docx.Document()
doc.add_heading("LỊCH CÔNG TÁC TUẦN", level=1)
doc.add_paragraph("Tuần từ 22/06/2026 đến 28/06/2026")
table = doc.add_table(rows=2, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Thứ/Ngày"
hdr_cells[1].text = "Thời gian"
hdr_cells[2].text = "Nội dung công việc"
hdr_cells[3].text = "Chủ trì"
hdr_cells[4].text = "Địa điểm"

row_cells = table.rows[1].cells
row_cells[0].text = "Thứ Hai\n22/06"
row_cells[1].text = "08:00 - 11:30"
row_cells[2].text = "Họp Giao ban đầu tuần phòng Quản lý đào tạo (QLĐT)"
row_cells[3].text = "Đ/c Vũ (PGĐ)"
row_cells[4].text = "Phòng hội trường lớn"

temp_doc_path = "temp_prod_test.docx"
doc.save(temp_doc_path)

ctx = ssl._create_unverified_context()

try:
    # 1. Login to production to get access token
    print("Logging in to production...")
    login_url = "https://qllichc500-1.onrender.com/api/auth/login"
    login_data = json.dumps({"username": "thupv", "password": "123456"}).encode('utf-8')

    req = urllib.request.Request(
        login_url,
        data=login_data,
        headers={"Content-Type": "application/json"}
    )
    with urllib.request.urlopen(req, context=ctx) as response:
        login_res = json.loads(response.read().decode('utf-8'))
        token = login_res["access_token"]
        print("Login successful. Access token obtained.")
        
    # 2. Upload file to production import API
    print("Uploading file to production import API...")
    import_url = "https://qllichc500-1.onrender.com/api/schedules/import"
    
    with open(temp_doc_path, "rb") as f:
        file_content = f.read()
        
    # Construct multipart request body manually
    boundary = b"----WebKitFormBoundary7MA4YWxkTrZu0gW"
    body = []
    body.append(b"--" + boundary)
    body.append(b'Content-Disposition: form-data; name="file"; filename="test.docx"')
    body.append(b'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    body.append(b'')
    body.append(file_content)
    body.append(b"--" + boundary + b"--")
    body.append(b'')
    
    encoded_body = b"\r\n".join(body)
    
    req_upload = urllib.request.Request(
        import_url,
        data=encoded_body,
        headers={
            "Content-Type": f"multipart/form-data; boundary={boundary.decode('utf-8')}",
            "Authorization": f"Bearer {token}"
        }
    )
    
    try:
        with urllib.request.urlopen(req_upload, context=ctx) as upload_res:
            res_content = upload_res.read().decode('utf-8')
            print("Import response status: 200 OK")
            res_json = json.loads(res_content)
            # Write to file with UTF-8 encoding
            with open("import_result.json", "w", encoding="utf-8") as out_f:
                json.dump(res_json, out_f, indent=2, ensure_ascii=False)
            print("Response successfully written to import_result.json")
    except urllib.error.HTTPError as e:
        print("Upload HTTPError status:", e.code)
        print("Upload HTTPError body:", e.read().decode('utf-8'))
        
except Exception as e:
    print("Test failed with exception:", e)
finally:
    if os.path.exists(temp_doc_path):
        os.remove(temp_doc_path)
